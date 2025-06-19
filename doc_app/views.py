import os
from pyexpat.errors import messages
import uuid

from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

import time
import logging
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import Employees
import re
import shutil
from docx import Document
from docx.shared import Pt
from doc_app.utils.doc_generator import DocumentGenerator
from django.conf import settings
from django.http import JsonResponse
from django.shortcuts import redirect, render

from doc_app.utils.doc_generator import DocumentGenerator
import pythoncom
from functools import wraps
logger = logging.getLogger(__name__)

def login_required_custom(view_func):
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.session.get('employee_id'):
            return redirect('login')  # Redirect to login page if not logged in
        return view_func(request, *args, **kwargs)
    return wrapper


def batch_parameters(folder_parameters, batch_size=5):
    """
    Yield batches of (folder, parameter) combinations
    """
    flat_list = []
    for folder, param_list in folder_parameters.items():
        for param in param_list:
            flat_list.append((folder, param))

    for i in range(0, len(flat_list), batch_size):
        yield flat_list[i:i + batch_size]


def extract_sections(text):
    pattern = r"###\s*Folder:\s*(Folder\s*\d+)\s*-\s*Parameter:\s*(.*?)\n"
    matches = list(re.finditer(pattern, text))

    sections = []

    for idx, match in enumerate(matches):
        folder = match.group(1).strip()
        parameter = match.group(2).strip()
        start = match.end()

        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        content = text[start:end].strip()

        sections.append((folder, parameter, content))

    return sections


from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def clean_and_format_content(doc, content):
    """
    Cleans markdown-like syntax and applies:
    - Font: Calibri
    - Main headings: 14pt bold
    - Subheadings: 12pt bold
    - Body: 11pt
    - Alignment: Left
    - Line spacing: 1.15
    - Spacing before/after: 6pt
    Removes all ** markdown stars
    """
    paragraphs = content.strip().split('\n')
    for line in paragraphs:
        line = line.strip()
        if not line:
            continue

        # Remove all ** stars
        line = line.replace("**", "")
        line.replace("*",""),
        line.replace("***", "")
        line = line.replace("###","")
        line = line.replace("##", "")
        line = line.replace("#","")

        # Main Heading: 1. Title
        main_heading_match = re.match(r"^([0-9]+\..+)$", line)
        if main_heading_match:
            heading = main_heading_match.group(1)
            para = doc.add_paragraph()
            run = para.add_run(heading)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            continue

        # Bullet Points (•, *, -)
        if re.match(r"^[*-•]\s+", line):
            text = re.sub(r"^[*-•]\s+", '', line)
            para = doc.add_paragraph(text, style='List Bullet')
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            continue

        # Sub-bullet: ___*, •, _ etc.
        if re.match(r"^[_•]+[*\-]?\s*", line):
            text = re.sub(r"^[_•]+[*\-]?\s*", '', line)
            para = doc.add_paragraph(text, style='List Bullet 2')
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            continue

        # Subheadings: Recognize like "Clear and Concise:"
        if re.match(r"^[A-Z][\w\s]+:\s*$", line):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            continue

        # Normal body text
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

def split_combined_doc_and_create_final_files(combined_path, company_name, company_address, doc_no, effective_date, logo_path=None):
    from doc_app.utils.doc_generator import DocumentGenerator

    doc = Document(combined_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])

    sections = extract_sections(full_text)
    if not sections:
        print("❌ No valid sections found.")
        return

    generator = DocumentGenerator()

    for folder_name, param_name, content in sections:
        folder_number = ''.join(filter(str.isdigit, folder_name)).zfill(2)
        folder_path = os.path.join("media", "final_output", folder_number)
        os.makedirs(folder_path, exist_ok=True)

        file_name = param_name.replace(" ", "_").replace("/", "-") + ".docx"
        file_path = os.path.join(folder_path, file_name)

        new_doc = Document()
        style = new_doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        generator._create_header(new_doc, company_name, company_address, doc_no, effective_date, logo_path)

        title_para = new_doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(param_name)
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_para.paragraph_format.space_after = Pt(12)

        clean_and_format_content(new_doc, content)

        new_doc.save(file_path)
        print(f"✅ Saved: {file_path}")

    # Cleanup
    combined_pdf_path = os.path.splitext(combined_path)[0] + ".pdf"
    if os.path.exists(combined_path):
        os.remove(combined_path)
        print(f"🗑️ Deleted: {combined_path}")
    if os.path.exists(combined_pdf_path):
        os.remove(combined_pdf_path)
        print(f"🗑️ Deleted: {combined_pdf_path}")












# === Run this directly ===
if __name__ == "__main__":
    split_combined_doc_and_create_final_files(
        input_path="media/generated_docs/31003_Combined_20250617_202328.docx",  # <- change if needed
        company_name="TCS",
        company_address="Vijay Nagar",
        doc_no="M.122.NC",
        effective_date="2025-06-17",
        logo_path="media/logos/logo.png"  # or None
    )


def login(request):
    return render(request, 'doc_app/login.html')



def verify_user(request):
    if request.method == 'POST':
        username = request.POST.get('username')  
        password = request.POST.get('password')

        try:
            user = Employees.objects.get(UserName=username)
            if user.Password == password:
                # Store user ID or name in session
                request.session['employee_id'] = user.id
                request.session['employee_name'] = user.UserName
                return redirect('division-selection')  
            else:
                messages.error(request, 'Invalid password')
                return redirect('login')  # Change to your login page name
        except Employees.DoesNotExist:
            messages.error(request, 'User not found')
            return redirect('login')

    return render(request, 'login.html')



def logout(request):
    request.session.flush()
    return redirect('login')






def home(request):
    return render(request, 'doc_app/welcome.html')

@login_required_custom
def division_selection(request):
    return render(request, 'doc_app/Nic_list.html')

@login_required_custom
def generate_documents(request):
    pythoncom.CoInitialize()  # For Windows COM objects

    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=400)

    try:
        data = {
            'company_name': request.POST.get('company-name', ''),
            'company_address': request.POST.get('company-address', ''),
            'effective_date': request.POST.get('EffDate', ''),
            'doc_no': request.POST.get('DocNO', ''),
            'nic_code': request.POST.get('nic_code', ''),
            'division_name': request.POST.get('division_name', '')
        }

        # Handle logo upload
        logo_path = None
        uploaded_logo = request.FILES.get('logo')
        if uploaded_logo:
            logo_folder = os.path.join(settings.MEDIA_ROOT, 'logos')
            os.makedirs(logo_folder, exist_ok=True)
            logo_filename = f"{uuid.uuid4()}_{uploaded_logo.name}"
            logo_path = os.path.join(logo_folder, logo_filename)
            with open(logo_path, 'wb+') as dest:
                for chunk in uploaded_logo.chunks():
                    dest.write(chunk)

        # Initialize generator
        template_dir = os.path.join('doc_app', 'templates', 'doc_app', 'document_templates')
        
        generator = DocumentGenerator(template_dir, gemini_api_key=settings.GEMINI_API_KEY)


        # Define folder -> parameters
        folder_parameters = {
            "Folder 1": [
                "Proper Nomenclature and Labelling",
                "Task Identification",
                "Employee Role Assignment",
                "Training of Employees",
                "Documentation and SOPs",
                "Work Area Layout",
                "Periodic Audits and Checks",
                "Root Cause Analysis and CAPA",
                "Implementation of 5S"
            ],
            "Folder 2": [
                "Establishing Safety Policy & Communicating to All Stakeholders",
                "Training of All Employees on Safety and Safety Requirements",
                "Providing Relevant Personal Protective Equipment (PPE)",
                "Signage & Pictograms/Pictures",
                "Maintaining Accident/Incident Register",
                "Inventory Knowledge and Classification",
                "Color-Coded Piping",
                "Spill Management System",
                "Insulation & Warning Signage",
                "Ventilation & Off-Gas Management",
                "Periodic Hydraulic Testing",
                "Provision of Worker Amenities",
                "Procedures for Non-Routine Operations",
                "Timely Availability of Medical/Emergency Services",
                "Maintenance of Emergency Response Equipment",
                "Mock Drills for Emergency Evacuations",
                "Periodic Safety Audits and CAPA"
            ],
            "Folder 3": [
                "Recording Contract Delivery Schedules",
                "Monitoring & Measuring Delivery Timelines (On-Time in Full)",
                "Use of Technology/Systems to Track and Measure Deliveries",
                "Additional Activities to Enhance Delivery Performance",
                "Example Workflow for Implementing Delivery Performance Systems"
            ],
            "Folder 4": [
                "Establishing Quality Policy & Communicating to All Stakeholders",
                "Identifying Quality Requirements",
                "Implementing Systems to Fulfil Quality Requirements",
                "Training All Employees on Quality & Processes",
                "Ensuring Employee Involvement in Quality Improvement Initiatives",
                "Periodic Quality Audits and Action Plans",
                "Additional Quality Management Activities",
                "Example Workflow for Implementing Quality Management Systems"
            ],
            "Folder 5": [
                "Defining Process Flow Diagrams (PFD) and Descriptions",
                "Tagging for Identification",
                "Identifying Critical Processes",
                "Sensors and Monitoring Systems",
                "Periodic Quality Audits and CAPA",
                "Additional Activities for Effective Process Control",
                "Example Workflow for Implementing Process Control Requirements"
            ],
            "Folder 6": [
                "Setting Up Targets for Quality, Cost & Delivery (QCD)",
                "Communicate the QCD Targets to Relevant Employees",
                "Display QCD Targets & Trends at Prominent Locations",
                "Monitor the Progress of QCD on a Daily Basis",
                "Address Deviations, Identify Root Causes, and Take Corrective & Preventive Actions (CAPA)",
                "Additional Activities for Effective Daily Work Management",
                "Example Workflow for Implementing a Daily Work Management System"
            ],
            "Folder 7": [
                "Scheduling Maintenance & Calibration",
                "Adherence to the Schedule",
                "Measuring Mean Time To Repair (MTTR) & Mean Time Before Failure (MTBF)",
                "Additional Activities for Effective Maintenance Management",
                "Example Workflow for Implementing a Maintenance Management System"
            ],
            "Folder 8": [
                "Defining Process Flow Diagrams (PFD) and Descriptions",
                "Tagging for Identification",
                "Identifying Critical Processes",
                "Sensors and Monitoring Systems",
                "Periodic Quality Audits and CAPA",
                "Additional Activities for Effective Process Control",
                "Example Workflow for Implementing Process Control Requirements"
            ],
            "Folder 9": [
                "Identifying Testing and Certification Requirements",
                "Acceptance Sampling for Product Quality",
                "Checking Products Against Specifications",
                "Provision of Laboratory Facility",
                "Documentation and Traceability",
                "Periodic Audits and CAPA",
                "Example Workflow for Implementing Testing and Certification Requirements"
            ],
            "Folder 10": [
                "Listing & Classification of Materials",
                "Employee Training",
                "Material Safety Data Sheets (MSDS)",
                "Storage Facility Management",
                "Scrap Material Management",
                "Material Handling & Transportation",
                "Inventory Control",
                "Warehouse Layout",
                "Periodic Audits and CAPA"
            ],
            "Folder 11": [
                "Identifying Environmental Aspects & Impacts",
                "Regulatory Compliance",
                "Infrastructure for Pollution Control",
                "Drainage Management",
                "Spill Management",
                "Proper Disposal of Paints and Thinners",
                "Training and Awareness",
                "Periodic Audits and CAPA"
            ],
            "Folder 12": [
                "Identifying Environmental Aspects & Impacts",
                "Regulatory Compliance",
                "Infrastructure for Pollution Control",
                "Drainage Management",
                "Spill Management",
                "Proper Disposal of Paints and Thinners",
                "Training and Awareness",
                "Periodic Audits and CAPA"
            ],
            "Folder 13": [
                "Defects",
                "Rework",
                "Rejection",
                "Cost of Poor Quality (COPQ)",
                "Customer Satisfaction",
                "Process for Measurement and Analysis"
            ],
            "Folder 14": [
                "Developing Processes to Identify, Select, Evaluate & Develop Suppliers",
                "Monitoring Performance of Vendors/Suppliers and Out-Sourced Partners Periodically",
                "Timely Communication to Vendors/Suppliers for Improvement and Sensitizing Them on Key Issues",
                "Summary"
            ],
            "Folder 15": [
                "Identifying All Risks",
                "Developing Plans to Address Risks",
                "Implementing Measures to Mitigate/Reduce Risks",
                "Reviewing Risks and Mitigation Plans Periodically"
            ],
            "Folder 16": [
                "Training of Employees on Muda, Mura, Muri",
                "Identifying Muda, Mura & Muri Wastes",
                "Establishing Targets & Action Plans to Eliminate/Reduce Wastes",
                "Monitoring of Muda, Mura, Muri Wastes"
            ],
            "Folder 17": [
                "Periodic Assessment of Existing Technology",
                "Training/Exposure to the Best Available/Latest Technology",
                "Planning and Adoption of Latest Technology",
                "Key Technologies to Consider",
                "Continuous Improvement and Monitoring"
            ],
            "Folder 18": [
                "Identification of Natural Resources Consumed",
                "Assessment of Possibility to Reduce Consumption & Recovery",
                "Establishing Targets for Natural Resource Conservation",
                "Workplace Design for Natural Light Utilization",
                "Ventilation/HVAC Design for Natural Air Draft",
                "Water Conservation Measures",
                "Training Relevant Employees & People",
                "Periodic Audits and CAPA"
            ],
            "Folder 19": [
                "Identifying Statutory & Regulatory Requirements",
                "Identifying Areas for Societal Contribution",
                "Establishing Targets & Resources",
                "Monitoring & Reviewing Progress",
                "Specific Areas for Societal Contribution",
                "Training & Awareness"
            ]
        }


        # Collect all content
        combined_sections = []

        for batch in batch_parameters(folder_parameters, batch_size=5):  # Feel free to increase or decrease batch size
            prompt_lines = []
            for folder, param in batch:
                prompt_lines.append(f"### Folder: {folder} - Parameter: {param}")

            batch_prompt = f"""
            You are a documentation expert.

            NIC Code: {data['nic_code']} – {data['division_name']}

            Generate content for the following sections. For each:

            - Write a heading like: ### Folder: <folder> - Parameter: <parameter>
            - Then generate 1–2 pages of professional documentation.
            - Include steps, compliance notes, and practical guidelines.
            - Format using Markdown-style sections.

            Sections:
            {chr(10).join(prompt_lines)}
            """

            # Get response from Gemini
            response = generator.get_gemini_response(prompt=batch_prompt)

            # Extract individual sections from the response
            import re
            matches = re.split(r"(### Folder: .*? - Parameter: .*?\n)", response)

            for i in range(1, len(matches), 2):
                section_title = matches[i].strip()
                section_body = matches[i+1].strip() if (i + 1 < len(matches)) else ""
                combined_sections.append(f"{section_title}\n\n{section_body}")


        # Combine into one DOCX
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        combined_doc_name = f"{data['nic_code']}_Combined_{timestamp}.docx"
        combined_doc_path = os.path.join(settings.MEDIA_ROOT, 'generated_docs', combined_doc_name)

        os.makedirs(os.path.dirname(combined_doc_path), exist_ok=True)

        generator.generate_combined_doc(
            output_path=combined_doc_path,
            sections=combined_sections,
            company_name=data['company_name'],
            company_address=data['company_address'],
            doc_no=data['doc_no'],
            effective_date=data['effective_date'],
            logo_path=logo_path
        )

        split_combined_doc_and_create_final_files(
            combined_path=combined_doc_path,
            company_name=data['company_name'],
            company_address=data['company_address'],
            doc_no=data['doc_no'],
            effective_date=data['effective_date'],
            logo_path=logo_path
        )

        try:
            os.remove(combined_doc_path)
            pdf_path = combined_doc_path.replace('.docx', '.pdf')
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
        except Exception as e:
            logger.warning(f"Failed to delete combined file: {e}")

        return render(request, 'doc_app/success.html')

    except Exception as e:
        logger.exception("Error generating combined document")
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


