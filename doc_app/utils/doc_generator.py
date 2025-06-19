import os
import time
import logging
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL  # ✅ Add this

import requests
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert

from django.conf import settings

logger = logging.getLogger(__name__)

class DocumentGenerator:
    
    def __init__(self, output_dir=None, logo_dir=None, gemini_api_key=None):
        self.output_dir = output_dir or os.path.join(settings.MEDIA_ROOT, 'sample')
        self.logo_dir = logo_dir or os.path.join(settings.MEDIA_ROOT, 'logos')
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.logo_dir, exist_ok=True)

        self.gemini_api_key = settings.GEMINI_API_KEY

        

        self.master_prompt = (
            "You are an expert in industrial documentation and compliance writing. For the NIC code: **{nic_code}**, "
            "generate a comprehensive, detailed document (2+ pages) for:\n\n"
            "📁 Folder: {folder}\n📄 Parameter: {file}\n\n"
            "Include:\n"
            "1. Detailed procedures and steps\n"
            "2. Relevant regulations/standards\n"
            "3. Safety considerations\n"
            "4. Quality control measures\n"
            "5. Implementation guidelines\n"
            "6. [Placeholder diagrams where needed]\n\n"
            "Use formal, professional language suitable for corporate documentation."
        )


    def _get_logo_path(self, company_name=None, uploaded_logo=None):
        if uploaded_logo and os.path.exists(uploaded_logo):
            return uploaded_logo
        possible_logos = [
            f"{company_name.lower().replace(' ', '_')}.png",
            f"{company_name.lower().replace(' ', '_')}.jpg",
            "default_logo.png",
            "default_logo.jpg"
        ]
        for file in possible_logos:
            full_path = os.path.join(self.logo_dir, file)
            if os.path.exists(full_path):
                return full_path
        return None

    def _create_header(self, doc, company_name, company_address, doc_no, effective_date, logo_path=None):
        section = doc.sections[0]
        table = doc.add_table(rows=2, cols=3)
        table.autofit = False
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set column widths
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(3.5)
        table.columns[2].width = Inches(2.5)

        # Column 1: Logo
        logo_cell = table.cell(0, 0)
        logo_cell.merge(table.cell(1, 0))
        logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo_path and os.path.exists(logo_path):
            try:
                logo_para.add_run().add_picture(logo_path, width=Inches(1.2))
            except Exception as e:
                logger.warning(f"Logo load error: {e}")
                logo_para.add_run("Logo Error")
        else:
            logo_para.add_run("No Logo")

        # Column 2: Company Info
        middle_cell = table.cell(0, 1)
        middle_cell.merge(table.cell(1, 1))
        middle_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        middle_para = middle_cell.paragraphs[0]
        middle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_name = middle_para.add_run(f"{company_name}\n")
        run_name.bold = True
        run_name.font.size = Pt(13)
        run_addr = middle_para.add_run(company_address)
        run_addr.font.size = Pt(11)

        # Column 3: DOC.NO and DATE
        doc_cell = table.cell(0, 2)
        eff_cell = table.cell(1, 2)
        doc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        eff_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc_para = doc_cell.paragraphs[0]
        doc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_doc = doc_para.add_run(f"DOC.NO: {doc_no}")
        run_doc.bold = True
        run_doc.font.size = Pt(11)

        eff_para = eff_cell.paragraphs[0]
        eff_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_eff_word = eff_para.add_run("EFFECTIVE\n")
        run_eff_word.bold = True
        run_eff_word.font.size = Pt(11)
        formatted_date = datetime.strptime(effective_date, "%Y-%m-%d").strftime("%d/%m/%Y")
        run_eff_date = eff_para.add_run(f"DATE: {formatted_date}")
        run_eff_date.bold = True
        run_eff_date.font.size = Pt(11)

        self.set_table_border(table)


    def set_table_border(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders')) or OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            element = tblBorders.find(qn(f'w:{border_name}')) or OxmlElement(f'w:{border_name}')
            tblBorders.append(element)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')

    def generate_docx(self, company_name, doc_no, effective_date, output_name=None, logo_path=None, company_address=""):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        self._create_header(doc, company_name, company_address, doc_no, effective_date, logo_path)
        doc.add_paragraph("Document content goes here...")

        if not output_name:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_name = f"{company_name.replace(' ', '_')}_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, output_name)
        doc.save(output_path)
        return output_path

    def generate_pdf(self, company_name, company_address, doc_no, effective_date, output_name=None, logo_path=None):
        docx_path = self.generate_docx(company_name, doc_no, effective_date, output_name, logo_path, company_address)
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
        convert(docx_path, pdf_path)
        return pdf_path
    
    
    def get_gemini_response(self, prompt: str) -> str:
        import google.generativeai as genai
        genai.configure(api_key=settings.GEMINI_API_KEY)
        model = genai.GenerativeModel("gemini-1.5-flash")  # or gemini-1.5-pro

        try:
            response = model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            logger.error(f"Gemini error: {e}")
            return f"[Error fetching content: {e}]"

    def generate_combined_doc(self, output_path, sections, company_name, company_address, doc_no, effective_date, logo_path=None):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        for section in sections:
            self._create_header(doc, company_name, company_address, doc_no, effective_date, logo_path)
            doc.add_paragraph(section)
            doc.add_page_break()

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
