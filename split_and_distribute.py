import os
import re
import shutil
from docx import Document
from docx.shared import Pt
from doc_app.utils.doc_generator import DocumentGenerator  # Adjust path if needed

def split_combined_doc_and_create_final_files(combined_path, company_name, company_address, doc_no, effective_date, logo_path=None):
    # Read combined DOCX
    doc = Document(combined_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])

    # Split using pattern
    sections = re.split(r"(### Folder: (.*?) - Parameter: (.*?)\n)", full_text)

    if not sections or len(sections) < 4:
        print("❌ No valid sections found.")
        return

    generator = DocumentGenerator()

    for i in range(1, len(sections), 4):  # Every 4th item: split header, folder, param, body
        full_heading = sections[i].strip()
        folder_name = sections[i + 1].strip()  # e.g., Folder 01
        param_name = sections[i + 2].strip()
        content = sections[i + 3].strip()

        folder_number = ''.join(filter(str.isdigit, folder_name)).zfill(2)
        folder_path = os.path.join("media", "final_output", folder_number)
        os.makedirs(folder_path, exist_ok=True)

        file_name = param_name.replace(" ", "_").replace("/", "-") + ".docx"
        file_path = os.path.join(folder_path, file_name)

        # Generate new document
        new_doc = Document()
        style = new_doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Add header and body
        generator._create_header(new_doc, company_name, company_address, doc_no, effective_date, logo_path)
        new_doc.add_paragraph(full_heading)
        for para in content.split('\n\n'):
            if para.strip():
                new_doc.add_paragraph(para.strip())

        new_doc.save(file_path)
        print(f"✅ Saved: {file_path}")

    # 🧹 Delete combined DOCX and PDF
    combined_pdf_path = os.path.splitext(combined_path)[0] + ".pdf"
    if os.path.exists(combined_path):
        os.remove(combined_path)
        print(f"🗑️ Deleted: {combined_path}")
    if os.path.exists(combined_pdf_path):
        os.remove(combined_pdf_path)
        print(f"🗑️ Deleted: {combined_pdf_path}")

# === Run this directly ===
if __name__ == "__main__":
    split_combined_doc_and_create_final_files(
        combined_path="media/generated_docs/31003_Combined_20250617_202328.docx",  # <- change if needed
        company_name="TCS",
        company_address="Vijay Nagar",
        doc_no="M.122.NC",
        effective_date="2025-06-17",
        logo_path="media/logos/logo.png"  # or None
    )
